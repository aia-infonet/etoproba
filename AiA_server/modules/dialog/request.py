"""
Команда "Запрос : начать запись" из модуля "Диалог".
Версия 3.2: стриминг + преобразование слов-цифр.
"""
import asyncio
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt
from typing import Dict, Any
import os

from utils.ollama import chat_stream_async

BASE_DIR = Path("D:/AiA") if os.path.exists("D:/") else Path.home() / "AiA"
DIALOG_DIR = BASE_DIR / "Dialog"

async def execute(ws_manager, client_ip: str, data: Dict[str, Any],
                  ollama_chat, ollama_model: str, tts_engine, logger):
    raw_text = data.get("text", "").strip()
    text = data.get("converted_text", raw_text).strip()

    if not raw_text:
        logger.warning("[request] Получен пустой текст от клиента")
        await ws_manager.send_to(client_ip, {
            "type": "error",
            "message": "Текст не распознан. Повторите, пожалуйста.",
            "module": "dialog",
            "command": "request"
        })
        return

    logger.info(f"[request] Запрос ({len(raw_text)} симв.): {raw_text[:200]}...")
    if text != raw_text:
        logger.info(f"[request] Преобразовано: '{text}'")

    try:
        now = datetime.now()
        year = str(now.year)
        filename = f"запрос_{now.month:02d}_{now.day:02d}_{now.hour:02d}_{now.minute:02d}.docx"

        year_dir = DIALOG_DIR / year
        year_dir.mkdir(parents=True, exist_ok=True)
        file_path = year_dir / filename

        doc = Document()
        title = doc.add_heading("Голосовой запрос и ответ", level=1)
        doc.add_paragraph(f"Дата: {now.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Источник: {client_ip}")
        doc.add_paragraph("")

        p_request = doc.add_paragraph()
        run_label = p_request.add_run("Запрос : ")
        run_label.bold = True
        run_label.font.size = Pt(12)
        p_request.add_run(text)

        logger.info("[request] Запрос записан в документ")

        await ws_manager.send_to(client_ip, {
            "type": "status",
            "message": "Обработка через LLM...",
            "module": "dialog",
            "command": "request"
        })

        messages = [{"role": "user", "content": raw_text}]
        full_text = ""
        chunk_count = 0

        async for chunk in chat_stream_async(ollama_model, messages, timeout=60):
            full_text += chunk
            chunk_count += 1
            await ws_manager.send_to(client_ip, {
                "type": "stream_chunk",
                "text": chunk,
                "module": "dialog",
                "command": "request"
            })

        logger.info(f"[request] Стрим завершен. Получено {chunk_count} чанков, итого {len(full_text)} симв.")

        await ws_manager.send_to(client_ip, {
            "type": "stream_end",
            "text": full_text,
            "module": "dialog",
            "command": "request"
        })

        if not full_text.strip():
            logger.error("[request] Полный текст пустой после стриминга! Озвучивание пропущено.")
            await ws_manager.send_to(client_ip, {
                "type": "error",
                "message": "Получен пустой ответ от модели.",
                "module": "dialog",
                "command": "request"
            })
            return

        logger.info(f"[request] Полный ответ LLM ({len(full_text)} симв.): {full_text[:200]}...")

        doc.add_paragraph("")
        p_response = doc.add_paragraph()
        run_resp_label = p_response.add_run("Ответ : ")
        run_resp_label.bold = True
        run_resp_label.font.size = Pt(12)
        p_response.add_run(full_text)

        doc.save(file_path)
        logger.info(f"[request] Файл {filename} сохранен в {year_dir}")

        await ws_manager.send_to(client_ip, {
            "type": "status",
            "message": f"Файл {filename} сохранен",
            "module": "dialog",
            "command": "request"
        })

    except asyncio.TimeoutError:
        logger.error("[request] Таймаут ожидания ответа от Ollama")
        await ws_manager.send_to(client_ip, {
            "type": "error",
            "message": "Модель слишком долго обрабатывает запрос. Попробуйте упростить вопрос.",
            "module": "dialog",
            "command": "request"
        })
    except Exception as e:
        logger.error(f"[request] Ошибка обработки: {e}", exc_info=True)
        await ws_manager.send_to(client_ip, {
            "type": "error",
            "message": f"Ошибка обработки: {e}",
            "module": "dialog",
            "command": "request"
        })
"""
Команда "Записки : начать запись" из модуля "Диалог".
Получает распознанный текст, создаёт .docx файл в D:\AiA\Dialog\YYYY\
с именем "запись_MM_DD_hh_mm.docx" и сохраняет в него текст.
"""
from datetime import datetime
from pathlib import Path
from docx import Document
from typing import Dict, Any
import os

BASE_DIR = Path("D:/AiA") if os.path.exists("D:/") else Path.home() / "AiA"
DIALOG_DIR = BASE_DIR / "Dialog"

async def execute(ws_manager, client_ip: str, data: Dict[str, Any],
                  ollama_chat, ollama_model: str, tts_engine, logger):
    raw_text = data.get("text", "").strip()
    text = data.get("converted_text", raw_text).strip()

    if not text:
        logger.warning("[record] Получен пустой текст")
        await ws_manager.send_to(client_ip, {
            "type": "error",
            "message": "Текст для записи пуст.",
            "module": "dialog",
            "command": "record"
        })
        return

    try:
        now = datetime.now()
        year = str(now.year)
        filename = f"запись_{now.month:02d}_{now.day:02d}_{now.hour:02d}_{now.minute:02d}.docx"

        year_dir = DIALOG_DIR / year
        year_dir.mkdir(parents=True, exist_ok=True)

        file_path = year_dir / filename

        doc = Document()
        doc.add_heading("Голосовая запись", level=1)
        doc.add_paragraph(f"Дата: {now.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Источник: {client_ip}")
        doc.add_paragraph("")
        doc.add_paragraph(text)
        doc.save(file_path)

        if text != raw_text:
            logger.info(f"Файл {filename} сохранён (исходный: '{raw_text}')")
        else:
            logger.info(f"Файл {filename} сохранён в {year_dir}")

        await ws_manager.send_to(client_ip, {
            "type": "status",
            "message": f"Запись сохранена: {filename}",
            "module": "dialog",
            "command": "record"
        })

    except Exception as e:
        logger.error(f"[record] Ошибка сохранения файла: {e}")
        await ws_manager.send_to(client_ip, {
            "type": "error",
            "message": f"Ошибка сохранения: {e}",
            "module": "dialog",
            "command": "record"
        })